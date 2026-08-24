from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"D:\capstone\docs\CS-30_完整流程与后续迭代计划.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "17365D"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
GREEN = "2F6B4F"
GOLD = "8A6500"
RED = "9B1C1C"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_rfonts(run, ascii_font="Calibri", east_asia="Microsoft YaHei"):
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia)


def set_style_font(style, ascii_font="Calibri", east_asia="Microsoft YaHei"):
    style.font.name = ascii_font
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_run(paragraph, text, bold=False, italic=False, color=None, size=None):
    run = paragraph.add_run(text)
    set_rfonts(run)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    return run


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        add_run(p, bold_prefix, bold=True, color=NAVY)
        add_run(p, text[len(bold_prefix):])
    else:
        add_run(p, text)
    return p


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    add_run(p, text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    add_run(p, text)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_rfonts(run)
    return p


def style_table_text(table, header=True, font_size=9.2):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.08
                if c_idx == 0 and len(row.cells) <= 2:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_rfonts(run)
                    run.font.size = Pt(font_size)
                    if header and r_idx == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor.from_string(NAVY)
            if header and r_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
    if header:
        set_repeat_table_header(table.rows[0])


def add_table(doc, headers, rows, widths_dxa, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, value in enumerate(headers):
        table.rows[0].cells[idx].text = value
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            row.cells[idx].text = str(value)
    set_table_geometry(table, widths_dxa)
    style_table_text(table, header=True, font_size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(doc, title, text, fill=CALLOUT, accent=NAVY):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_run(p, title, bold=True, color=accent, size=11)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    add_run(p2, text, size=10.5)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_repeat_table_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(paragraph, "CS-30 项目迭代计划  |  ", color=MUTED, size=9)
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_rfonts(run)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    set_style_font(normal)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    h1 = doc.styles["Heading 1"]
    set_style_font(h1)
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(BLUE)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    set_style_font(h2)
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(BLUE)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    set_style_font(h3)
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = doc.styles[style_name]
        set_style_font(style)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375 if style_name != "List Bullet 2" else 0.625)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_styles(doc)

    # Running header/footer: quiet implementation-plan furniture.
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_p.paragraph_format.space_after = Pt(0)
    add_run(header_p, "CS-30  |  Delivery & Iteration Plan", color=MUTED, size=9)
    add_page_number(section.footer.paragraphs[0])

    # Customer-pack style opening block.
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(3)
    add_run(kicker, "PROJECT DELIVERY PLAN", bold=True, color=BLUE, size=10)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    add_run(title, "CS-30 完整流程与后续迭代计划", bold=True, color=NAVY, size=25)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    add_run(
        subtitle,
        "基于 OpenStax Physics、SciQ 评估集、个性化 RAG 与 Restricted KG 的 14 周 DevOps 路线",
        color=MUTED,
        size=12,
    )

    metadata = add_table(
        doc,
        ["计划基线", "团队投入", "工作方式", "文档日期"],
        [["第1周保守版 Thin Slice", "8人 × 20小时/周", "一周Sprint + 持续集成", "2026-08-23"]],
        [2340, 2340, 2340, 2340],
        font_size=9.4,
    )
    for cell in metadata.rows[1].cells:
        set_cell_shading(cell, WHITE)

    add_callout(
        doc,
        "推荐结论",
        "按照当前保守的第一周进度，完整范围建议按 14 周规划。第 7 周形成不含 KG 的核心系统，"
        "第 9 周完成 Restricted KG Beta，第 10 周冻结正式实验配置，第 14 周完成报告、代码与最终演示。"
        "若不纳入 KG，可压缩至约 11–12 周。",
        fill=LIGHT_BLUE,
        accent=NAVY,
    )

    add_heading(doc, "1. 估算口径与范围", 1)
    add_body(
        doc,
        "本计划假设团队共有 8 人，每人每周投入约 20 小时。名义产能为 160 人时/周；考虑会议、客户汇报、"
        "代码 Review、环境配置、数据返工和实验失败后，有效交付产能按 120–135 人时/周估算。"
    )
    add_body(doc, "“完整流程”包含以下范围：")
    for item in [
        "OpenStax Physics 知识库解析、可追溯 Chunking、Metadata 与索引构建；",
        "SciQ Physics 题目筛选、Evidence Alignment、章节隔离的 Smoke / Dev / Test；",
        "BM25、Dense、RRF Hybrid Retrieval 与检索侧评估；",
        "Synthetic Student Profile、Evidence Role、Level-Aware Reranking 与 Personalised Prompt；",
        "Calibrated Abstention、Citation Integrity、答案质量与 groundedness 评价；",
        "Restricted KG、路由、邻居扩展与候选重新计分；",
        "主模型个性化实验、Llama / Qwen / Gemma / GPT 横向实验、人工盲评；",
        "代码、可复现环境、客户周报、最终报告与演示。",
    ]:
        add_bullet(doc, item)
    add_callout(
        doc,
        "范围边界",
        "本估算不包含真实成绩单入口的伦理审批，也不包含正式真人学习成效研究。主实验使用固定的 Synthetic Profiles。"
        "如果增加真实学生实验，通常还需要额外 4–8 周，并可能受伦理审批时间影响。",
        fill=CALLOUT,
        accent=GOLD,
    )

    add_heading(doc, "2. 总工期判断", 1)
    add_table(
        doc,
        ["交付层级", "预计周期", "完成内容"],
        [
            ["可演示 Thin Slice", "1周", "小规模 OpenStax → Dense → SciQ → LLM → Evaluation"],
            ["核心系统（不含KG）", "7周可实现；11–12周可完成实验与报告", "Hybrid、个性化、可靠性、正式实验和报告"],
            ["完整系统（含Restricted KG）", "13–14周", "核心系统 + KG Beta + KG专项实验 + 正式评价"],
            ["含真人研究或成绩单", "额外4–8周以上", "伦理、招募、授权、数据治理和真人评价"],
        ],
        [2200, 2100, 5060],
        font_size=9.5,
    )
    add_body(
        doc,
        "因此，团队对客户的合理承诺应是：第 1 周交付可运行 v0.1；第 7 周交付不含 KG 的核心系统；"
        "第 9 周交付 KG Beta；第 10 周进入正式实验冻结；第 14 周完成最终交付。"
    )

    add_heading(doc, "3. 迭代式工作逻辑", 1)
    add_body(
        doc,
        "项目不按照“数据全部完成后再写检索、检索完成后再写生成”的瀑布方式推进。每周选择一条可展示的"
        "垂直切片，使用 Mock / Fixture 数据并行开发，再替换为真实数据。"
    )
    add_callout(
        doc,
        "每周循环",
        "客户需求或研究假设 → Backlog 排序 → 本周 Vertical Slice → 开发与持续集成 → Staging Release → "
        "客户 Demo 与指标汇报 → 客户反馈进入下一周 Backlog。",
        fill=LIGHT_BLUE,
        accent=BLUE,
    )
    add_body(doc, "每周固定节奏：")
    for item in [
        "Sprint Planning：确定本周可演示版本、验收标准和需要客户决定的问题；",
        "持续开发：小批量 PR、自动测试、Smoke Set 回归、数据和索引版本化；",
        "Staging Release：本周功能合并到稳定演示环境，不在客户会议前临时拼装；",
        "客户 Review：展示系统、指标、案例、风险和下一周建议；",
        "Retro：分析阻塞和返工原因，重新排序 Backlog。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4. 14周迭代路线", 1)
    roadmap_rows = [
        ["第1周\nv0.1 Thin Slice", "2–3章 OpenStax；20–30道 Pilot；500-token Chunking；Dense FAISS；10–20道 JSON 回答。", "现场运行一题；确认接口和技术链路；决定第2周数据扩展重点。"],
        ["第2周\nv0.2 Data Expansion", "扩大教材解析和 SciQ Physics 筛选；提升 Alignment 数量；建立数据质量 Dashboard。", "展示题量、章节覆盖和 Alignment 失败原因；确认 Physics 是否继续。"],
        ["第3周\nv0.3 Benchmark Candidate", "稳定 parser、char span 和 gold evidence；准备 Smoke / Dev / Test 章节隔离方案。", "展示可追溯 gold evidence；决定是否具备 60 Dev + 180 Test 条件。"],
        ["第4周\nv0.4 Hybrid Retrieval", "300/500 chunk；embedding Pilot；BM25、Dense 和 RRF；纯检索调参和缓存。", "对比 Hit@K、Recall@K、MRR、延迟和失败案例；选择主检索配置。"],
        ["第5周\nv0.5 Evidence & Evaluation", "Evidence Role 试标、20%双标、IAA；完成 D1/D2；设计 D3/D4 评价表。", "展示 taxonomy 分布和一致性；决定是否冻结 Evidence Role。"],
        ["第6周\nv0.6 Personalisation", "Synthetic Profiles；C4 Level-Aware Reranking；C6 Personalised Prompt；独立开关与 λ 扫描。", "同一问题三档回答；展示检索个性化和 Prompt 个性化的独立作用。"],
        ["第7周\nv0.7 Reliability", "C5 Abstention、C8 Citation Integrity、D4/D5；核心非KG系统整合。", "展示回答、拒答和非法引用案例；验收非KG核心系统。"],
        ["第8周\nv0.8 KG Spike", "定义 Restricted KG schema；抽取概念与关系；所有节点链接回原始 chunk；建立路由原型。", "展示小规模概念图和证据追溯；KG Go / No-Go 决策。"],
        ["第9周\nv0.9 KG Beta", "Hybrid → KG Expansion → BM25/Dense重新计分 → RRF → C4；构造multi-hop子集。", "比较 Hybrid 与 Hybrid+KG；决定KG进入正式结果还是保留为扩展。"],
        ["第10周\nv1.0 Release Candidate", "全系统集成；Dev调参；Prompt、λ、Top-K、模型和缓存冻结；四模型小规模预跑。", "展示完整 feature flags；审批正式实验矩阵和预算。"],
        ["第11周\nMain Experiment", "主模型 180题 × 12 condition，共2160次调用；自动指标与异常重试。", "汇报主实验完成率、成本、初步个性化结果和异常输出。"],
        ["第12周\nModel Comparison", "另外3个模型 × 60题 × 6 condition，共1080次调用；汇总3240次结果。", "展示 Llama / Qwen / Gemma / GPT 横向比较和自动指标。"],
        ["第13周\nHuman Evaluation", "360份人工盲评；其中72份双标；IAA、统计分析、成功与失败案例；报告持续整合。", "展示人工评价、分歧、一致性和主要研究结论。"],
        ["第14周\nFinal Release", "结果复核、缺失实验重跑、报告定稿、代码清理、复现检查和最终演示。", "最终系统、代码、报告、演示和项目回顾。"],
    ]
    add_table(
        doc,
        ["周次 / 版本", "工程与研究增量", "客户汇报与决策"],
        roadmap_rows,
        [1700, 4200, 3460],
        font_size=8.8,
    )

    add_heading(doc, "5. 八人后续职责安排", 1)
    role_rows = [
        ["1号 Leader", "架构、接口、CI/CD、集成、客户沟通", "持续维护主流水线、feature flags、Release、架构与报告整合。"],
        ["2号", "OpenStax解析和数据版本", "数据稳定后负责重建自动化、解析回归、数据卡和Implementation写作。"],
        ["3号", "SciQ分类、清洗和章节映射", "数据稳定后负责实验题目采样、分层统计和人工评价样本管理。"],
        ["4号", "Evidence Alignment", "后续转向 KG 概念/关系证据链接、Unresolved Pool和D4 claim审核。"],
        ["5号", "Chunking、Metadata、Evidence Role", "后续负责A3标注、C4 Level-Aware Reranking和 λ 实验。"],
        ["6号", "BM25、Dense、RRF与性能", "后续负责C3 KG Expansion、候选重新计分、检索消融和缓存。"],
        ["7号", "Profile、Prompt和LLM", "后续负责C5 Abstention、C8 Citation、多模型适配和调用稳定性。"],
        ["8号", "Evaluation、QA和实验管理", "持续负责D1–D5、Smoke回归、正式实验完整性、人工盲评与统计。"],
    ]
    add_table(
        doc,
        ["成员", "持续主责", "后续职责变化"],
        role_rows,
        [1500, 3100, 4760],
        font_size=9.2,
    )

    add_heading(doc, "6. 防止人员联动阻塞", 1)
    add_body(
        doc,
        "系统存在天然依赖，但不应让下游成员等待上游完整交付。团队采用“接口先行、样例先行、Mock与真实数据双轨”的方式。"
    )
    dep_rows = [
        ["2号 → 5号", "教材未解析，无法Chunking", "Leader提供Document fixture；2号先交一个章节样例。"],
        ["2/3号 → 4号", "教材或题目不足，无法Alignment", "先用5–10道手工样例验证算法，再扩大批量。"],
        ["5号 → 6号", "没有真实chunks，无法索引", "6号先用模拟chunks开发FAISS，随后替换真实数据。"],
        ["6号 → 7号", "没有Top-K，无法开发Prompt", "使用固定RetrievalResult fixture开发生成模块。"],
        ["4/6号 → 8号", "没有gold和检索结果，无法计算指标", "8号用人工构造小样例先验证指标实现。"],
        ["全员 → Leader", "周末首次合并导致集成失败", "持续小批量PR；每周Staging前已有可运行主线。"],
    ]
    add_table(
        doc,
        ["依赖", "风险", "并行化措施"],
        dep_rows,
        [1500, 3000, 4860],
        font_size=9.2,
    )
    add_callout(
        doc,
        "执行原则",
        "每个模块先交约10%的可用样例让下游开工，再继续扩大到本周目标。周五不是第一次集成；"
        "周五只应是本周稳定版本的客户Review。",
        fill=CALLOUT,
        accent=GREEN,
    )

    add_heading(doc, "7. 研究冻结点", 1)
    add_body(
        doc,
        "DevOps允许持续改进系统，但正式研究需要可比性。以下内容到达冻结点后只能通过新版本和新实验批次修改："
    )
    freeze_rows = [
        ["第3周末", "教材版本、parser version、gold evidence规则和数据划分候选"],
        ["第4周末", "主chunk size、embedding、Top-K和RRF配置"],
        ["第5周末", "Evidence Role taxonomy和人工标注指南"],
        ["第10周末", "Prompt、λ、模型版本、温度、token budget、实验矩阵和人工抽样方案"],
        ["正式Test开启后", "禁止使用Test调参；任何变更必须新建实验版本并完整重跑受影响条件"],
    ]
    add_table(doc, ["时间", "冻结内容"], freeze_rows, [1900, 7460], font_size=9.5)

    add_heading(doc, "8. 正式实验与评价安排", 1)
    add_heading(doc, "8.1 自动实验", 2)
    for item in [
        "个性化主实验：主模型 × 180题 × 12 condition = 2160次调用；",
        "模型对比：另外3个模型 × 60题 × 6 condition = 1080次调用；",
        "总计3240次LLM调用，主模型在60题子集上的结果复用；",
        "全部答案计算Answer-choice Accuracy、可读性、citation integrity和检索指标；",
        "检索侧调参和缓存不调用LLM。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "8.2 人工评价", 2)
    for item in [
        "从主模型结果中分层抽取360份进行六维盲评、水平适配和atomic claim评价；",
        "其中72份进行双标，计算Cohen κ或Krippendorff α；",
        "8人平均约45份主标注，并交叉承担双标任务；",
        "人工评价身份、模型和condition必须隐藏。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "9. 每周客户汇报模板", 1)
    add_table(
        doc,
        ["模块", "汇报内容"],
        [
            ["本周承诺", "本周计划交付的Vertical Slice和验收标准。"],
            ["现场Demo", "使用Staging环境展示真实数据流，不使用临时手工拼接。"],
            ["指标变化", "与上周版本比较Hit@K、MRR、Accuracy、格式成功率、延迟和成本。"],
            ["案例分析", "至少一个成功案例和一个失败案例，说明错误属于哪个模块。"],
            ["风险与决策", "需要客户确认的数据范围、预算、KG地位、评价方法或优先级。"],
            ["下周建议", "列出2–3个可完成的版本目标，不承诺过多并行功能。"],
        ],
        [1900, 7460],
        font_size=9.5,
    )

    add_heading(doc, "10. 关键风险、触发条件与缓冲", 1)
    risk_rows = [
        ["SciQ Physics题量不足", "第2–3周候选题无法支持60 Dev + 180 Test", "比较Biology/Chemistry或缩小矩阵；可能增加1–2周"],
        ["Alignment质量不足", "人工抽验错误较多", "修订匹配规则、扩大人工审核；优先保证可信子集"],
        ["解析逻辑反复变化", "char span漂移、索引重复重建", "第3周设parser冻结点；变更必须提升版本号"],
        ["Evidence Role一致性不足", "comparison/application频繁混淆", "修订定义并重新试标；可能增加1周"],
        ["KG引入噪声", "Hybrid+KG不优于Hybrid", "KG保留为扩展结果，不阻塞主实验"],
        ["API/GPU不足", "模型调用或embedding无法按时完成", "优先本地小规模预跑；正式调用分批执行并缓存"],
        ["正式实验缺失", "API失败、JSON错误或condition缺失", "第14周缓冲重跑；每批运行后立即做完整性检查"],
    ]
    add_table(
        doc,
        ["风险", "触发条件", "处理与进度影响"],
        risk_rows,
        [2300, 3300, 3760],
        font_size=9.0,
    )

    add_heading(doc, "11. 完成定义", 1)
    add_body(doc, "项目只有同时满足以下条件才视为完成：")
    for item in [
        "代码已合并至主分支，全部主要功能可通过feature flags独立开启或关闭；",
        "OpenStax、SciQ、Alignment、chunks、索引、Prompt、模型和实验结果均有版本记录；",
        "正式Test未被用于调参，实验矩阵没有缺失condition；",
        "自动指标、人工评价、IAA、成功与失败案例均已完成；",
        "Restricted KG有独立消融结果，或有清晰证据说明为何保留为扩展；",
        "另一名成员可根据README在干净环境中复现主要流程；",
        "客户已完成最终Review，报告、代码和演示材料一致。",
    ]:
        add_bullet(doc, item)

    add_callout(
        doc,
        "最终建议",
        "将 14 周作为完整含 KG 版本的基准计划，并把第 14 周作为明确缓冲，不在该周新增功能。"
        "每周只承诺一个可运行的垂直增量；如果 KG 或题量出现问题，优先保护第 7 周完成的非 KG 核心系统和正式实验有效性。",
        fill=LIGHT_BLUE,
        accent=NAVY,
    )

    # Keep table rows intact where possible, but allow natural expansion and page splitting.
    for table in doc.tables:
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            cant_split = OxmlElement("w:cantSplit")
            tr_pr.append(cant_split)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_document()
    print(path)
