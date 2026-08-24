from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from build_complete_project_plan_docx import (
    BLUE,
    DARK_BLUE,
    LIGHT_BLUE,
    LIGHT_GRAY,
    MUTED,
    NAVY,
    WHITE,
    add_body,
    add_bullet,
    add_callout,
    add_heading,
    add_page_number,
    add_run,
    add_table,
    configure_styles,
    set_cell_shading,
)


OUTPUT = Path(r"D:\capstone\docs\CS-30_第一周分工安排_不含正式评估.docx")


def add_bullet(doc, text, level=0):
    """Add a renderer-stable bullet instead of relying on Word list numbering."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.28 + 0.18 * level)
    p.paragraph_format.first_line_indent = Inches(-0.16)
    add_run(p, "• ", bold=True)
    add_run(p, text)
    return p


def add_role_section(doc, number, role, responsibility, work_items, deliverables, completion):
    add_heading(doc, f"{number}号：{role}", 1)

    add_heading(doc, "负责方向", 2)
    add_body(doc, responsibility)

    add_heading(doc, "本周具体工作", 2)
    for item in work_items:
        add_bullet(doc, item)

    add_heading(doc, "最终交付物", 2)
    for item in deliverables:
        add_bullet(doc, item)

    add_heading(doc, "完成标准", 2)
    for item in completion:
        add_bullet(doc, item)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1.15)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    add_run(header, "CS-30  |  Week 1 Delivery Plan", color=MUTED, size=9)
    add_page_number(section.footer.paragraphs[0])

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(3)
    add_run(kicker, "WEEK 1 DELIVERY PLAN", bold=True, color=BLUE, size=10)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    add_run(title, "CS-30 第一周分工安排", bold=True, color=NAVY, size=25)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    add_run(
        subtitle,
        "不含正式评估｜OpenStax Physics + Dense RAG Thin Slice",
        color=MUTED,
        size=12,
    )

    metadata = add_table(
        doc,
        ["团队规模", "投入时间", "迭代方式", "版本目标"],
        [["8人", "每人约20小时", "一周Sprint + DevOps", "v0.1-thin-slice"]],
        [1800, 2200, 2860, 2500],
        font_size=9.5,
    )
    for cell in metadata.rows[1].cells:
        set_cell_shading(cell, WHITE)

    add_callout(
        doc,
        "本周目标",
        "跑通 OpenStax Physics 2–3个章节 → 500-token Chunking → Embedding → FAISS Dense Retrieval → "
        "Student Profile → Personalised Prompt → LLM固定JSON回答 → 简单演示界面。"
        "第一周只验证技术链路，不输出正式研究评估结论。",
        fill=LIGHT_BLUE,
        accent=NAVY,
    )

    add_heading(doc, "1. 第一周范围", 1)
    add_heading(doc, "1.1 必须完成", 2)
    must_items = [
        "可靠解析 OpenStax Physics 的2–3个章节；",
        "从 SciQ 中整理20–30道Physics演示问题；",
        "生成一套带章节和来源信息的500-token chunks；",
        "使用一个embedding模型建立FAISS Dense索引；",
        "输入问题后能够返回Top-K chunks和教材来源；",
        "支持Beginner、Intermediate、Advanced三档学生水平；",
        "至少10–20道问题可以生成固定JSON回答；",
        "回答中的citation ID必须来自实际检索结果；",
        "提供统一运行入口和简单演示界面；",
        "周五可以在Staging环境向客户现场演示。",
    ]
    for item in must_items:
        add_bullet(doc, item)

    add_heading(doc, "1.2 本周不做", 2)
    not_do = [
        "Evidence Alignment正式流程和gold char span标注；",
        "Hit@K、Recall@K、MRR和Answer-choice Accuracy；",
        "正式Dev/Test划分和正式Test实验；",
        "Evidence Role标注、BM25/RRF正式比较和Abstention校准；",
        "Restricted KG和多模型正式对比。",
    ]
    for item in not_do:
        add_bullet(doc, item)

    add_callout(
        doc,
        "工程检查不等于正式评估",
        "本周仍需检查系统能否启动、数据能否读取、索引能否加载、API能否返回结果、JSON能否解析、"
        "citation ID是否存在，以及另一名成员能否按照README运行。这些检查只证明系统可运行，"
        "不证明某个模型或检索方案效果更好。",
        fill=LIGHT_GRAY,
        accent=DARK_BLUE,
    )

    add_heading(doc, "2. 八人分工总览", 1)
    overview_rows = [
        ["1号 Leader", "架构、接口、集成、客户沟通", "端到端主流水线和Staging Demo"],
        ["2号", "OpenStax教材解析与清洗", "2–3章标准化语料和解析程序"],
        ["3号", "SciQ与演示问题准备", "20–30道Physics演示题和自由问答问题"],
        ["4号", "Chunking与Metadata", "500-token chunks和原文追溯"],
        ["5号", "Embedding与FAISS索引", "Dense索引、映射文件和加载功能"],
        ["6号", "Dense Retrieval与Backend API", "Top-K Retriever和统一API"],
        ["7号", "Profile、Prompt与LLM", "三档画像、固定JSON和生成结果"],
        ["8号", "平台工程、演示界面与系统QA", "可复现环境、界面、Smoke Test和README"],
    ]
    add_table(
        doc,
        ["成员", "负责方向", "本周核心交付"],
        overview_rows,
        [1500, 3500, 4360],
        font_size=9.2,
    )

    add_role_section(
        doc,
        1,
        "Leader／系统架构与集成",
        "负责系统架构、模块接口、主流水线、持续集成和客户沟通。Leader需要参与架构与集成代码，不能只负责会议和任务分配。",
        [
            "编写系统组件图、数据流图和端到端调用流程；",
            "定义并冻结OpenStaxDocument、Chunk、SciQQuestion、RetrievalResult、StudentProfile和GeneratedAnswer接口；",
            "确定document_id、chapter_id、chunk_id、source、char_start和char_end等公共字段；",
            "搭建统一项目目录、配置、日志、错误处理和运行入口；",
            "为未完成模块提供Mock / Fixture数据，保证下游可以并行开发；",
            "持续集成2–8号的模块，Review跨模块代码并解决接口冲突；",
            "准备周五客户Demo、周报、风险清单和下周Backlog。",
        ],
        [
            "系统组件图、数据流图和接口说明；",
            "Architecture Decision Record；",
            "可运行的端到端主流水线；",
            "统一配置、日志和启动入口；",
            "Staging Demo、客户周报和下周Backlog。",
        ],
        [
            "所有模块使用同一套接口；",
            "至少一条真实问题能够经过完整流水线；",
            "Demo不依赖手工修改路径或临时Notebook；",
            "另一名成员可以依据README复现运行。",
        ],
    )

    add_role_section(
        doc,
        2,
        "OpenStax数据工程",
        "负责OpenStax Physics教材获取、版本记录、文本解析、清洗、章节化和原文位置保留。",
        [
            "确定并获取一本OpenStax Physics教材，记录名称、版本、来源、下载日期和文件格式；",
            "计算并记录document_hash，维护parser_version；",
            "选择2–3个适合第一周演示的章节；",
            "解析章节、标题、小节、正文和页码；",
            "清理页眉、页脚、重复内容、乱码和明显断词；",
            "保留稳定文本位置，支持chunk返回教材原文；",
            "人工检查至少10个段落或小节，记录公式、表格、图片和跨页文本问题；",
            "先交一个可用章节给4号，再继续完成其他章节。",
        ],
        [
            "OpenStax Physics原始教材及版本记录；",
            "2–3个章节的标准化语料；",
            "教材解析程序、document_hash和parser_version；",
            "解析质量抽查记录和已知问题清单。",
        ],
        [
            "相同输入能够重复生成相同标准化文本；",
            "章节、标题和正文没有明显错位；",
            "标准化文本可被4号直接用于Chunking；",
            "任意演示chunk可以定位回教材原文。",
        ],
    )

    add_role_section(
        doc,
        3,
        "SciQ与演示问题准备",
        "负责SciQ数据读取、Physics问题筛选和客户演示问题准备。本周不构建正式评估集。",
        [
            "读取并标准化SciQ数据，保留问题、四个选项、正确答案和support；",
            "筛选20–30道Physics问题，并确保与试验章节大致相关；",
            "检查选项顺序和正确答案字段，标记明显超出章节范围的问题；",
            "准备5–10道适合客户现场输入的自由问答问题；",
            "与7号确认问题格式能够直接进入Prompt；",
            "明确本周题目仅用于Smoke和Demo，不作为正式Test。",
        ],
        [
            "20–30道标准化SciQ Physics演示题；",
            "SciQ读取与转换程序；",
            "统一问题、选项、答案和support格式；",
            "5–10道自由问答演示问题；",
            "章节范围外问题清单和Demo题目说明。",
        ],
        [
            "每道选择题都有完整问题、四个选项和答案字段；",
            "问题格式能够被6号和7号直接读取；",
            "演示问题覆盖定义、解释和简单应用；",
            "没有将演示题错误称为正式评估集。",
        ],
    )

    add_role_section(
        doc,
        4,
        "Chunking与Metadata",
        "负责将标准化OpenStax文本转换成可检索chunks，并保留教材来源、章节和原文追溯信息。",
        [
            "实现500-token Chunking并设计合理overlap；",
            "为每个chunk生成唯一chunk_id；",
            "保留document_id、chapter_id、source、正文和char span；",
            "写入基础chapter/topic metadata；",
            "检查空chunk、重复chunk、极短chunk和跨章节切分；",
            "与2号共同验证至少10个chunk的原文追溯；",
            "先提供20–50个样例chunks给5号开发索引；",
            "主版本稳定前不做300-token对照实验。",
        ],
        [
            "500-token Chunking程序和配置；",
            "2–3个章节的完整chunks及metadata；",
            "Chunk数量、长度和章节分布统计；",
            "原文追溯样例和已知问题清单。",
        ],
        [
            "所有chunk都有唯一ID和章节来源；",
            "chunk文本可以定位回标准化教材；",
            "不存在明显空chunk和跨章节混合；",
            "输出可被5号直接用于embedding和索引构建。",
        ],
    )

    add_role_section(
        doc,
        5,
        "Embedding与FAISS索引",
        "负责embedding模型接入、chunk向量生成、FAISS索引构建、持久化和版本记录。",
        [
            "选择一个适合第一周Pilot的embedding模型，检查查询侧指令前缀；",
            "为chunks批量生成embeddings并执行一致的向量归一化；",
            "使用FAISS IndexFlatIP建立精确索引；",
            "建立chunk_id与向量位置的映射；",
            "实现索引保存、加载和版本记录；",
            "记录模型、维度、运行设备和构建时间；",
            "先使用样例chunks建立小索引，再替换为真实章节数据；",
            "与6号确认索引调用接口。",
        ],
        [
            "Embedding生成程序和FAISS索引构建程序；",
            "2–3个章节的Dense索引；",
            "chunk_id映射文件；",
            "索引保存与加载功能；",
            "模型、配置、维度和构建时间记录；",
            "索引使用说明。",
        ],
        [
            "相同语料和配置可以重复构建索引；",
            "保存后的索引能够重新加载；",
            "加载后能够返回对应chunk_id；",
            "6号可以通过固定接口使用索引。",
        ],
    )

    add_role_section(
        doc,
        6,
        "Dense Retrieval与Backend API",
        "负责将用户问题转换为查询向量、检索Top-K chunks，并通过统一Backend API提供给生成模块。",
        [
            "实现问题预处理和查询embedding；",
            "调用5号的FAISS索引并实现Top-K Dense Retrieval；",
            "返回chunk_id、文本、章节、来源、分数和排名；",
            "封装Retrieval API或统一Python service；",
            "实现索引缺失、空问题和无结果时的错误处理；",
            "缓存本周演示问题的检索结果；",
            "准备至少两个检索结果合理的演示案例；",
            "与7号确认RetrievalResult接口；",
            "真实索引完成前使用模拟索引或固定chunks开发API。",
        ],
        [
            "Top-K Dense Retriever和Backend API；",
            "API输入输出说明；",
            "检索结果缓存、错误处理和日志；",
            "真实演示问题的Top-K输出样例；",
            "与生成模块的连接代码。",
        ],
        [
            "输入问题后能够稳定返回Top-K chunks；",
            "结果包含教材来源和chunk ID；",
            "7号可以直接使用结果构建Prompt；",
            "异常时返回明确错误而不使系统退出。",
        ],
    )

    add_role_section(
        doc,
        7,
        "Student Profile、Prompt与LLM Generation",
        "负责学生画像、个性化Prompt、LLM接入、固定JSON输出和citation完整性。",
        [
            "接入一个可用LLM并建立最小StudentProfile schema；",
            "支持Beginner、Intermediate和Advanced三档水平；",
            "实现问题、Student Profile和Top-K evidence的Prompt组装；",
            "固定final_choice、explanation和citations三个JSON字段；",
            "实现JSON解析、schema校验和有限重试；",
            "限制citations只能引用输入的chunk ID；",
            "处理API超时、空响应和非法JSON；",
            "完成至少10–20道问题的端到端生成；",
            "准备同一道题在三个水平下的回答样例；",
            "记录模型、温度、token使用量和失败类型；",
            "6号完成前使用固定RetrievalResult fixture开发。",
        ],
        [
            "LLM生成适配器和Student Profile schema；",
            "Beginner、Intermediate和Advanced三档配置；",
            "Prompt模板、固定JSON schema和解析器；",
            "至少10–20道生成结果；",
            "三档回答样例及citation、失败和token记录。",
        ],
        [
            "输出能够稳定解析为固定JSON；",
            "citation ID来自实际检索输入；",
            "三档水平能够进入Prompt；",
            "单次API失败不会导致整个批次终止。",
        ],
    )

    add_role_section(
        doc,
        8,
        "平台工程、演示界面与系统QA",
        "负责将团队模块整理成可复现、可启动、可操作和可向客户演示的系统，不进行正式研究评估。",
        [
            "整理Python版本、依赖、环境变量和API Key配置；",
            "创建依赖锁定文件或明确依赖清单，提供.env.example且不提交真实密钥；",
            "编写一键启动或统一启动脚本；",
            "搭建简单Streamlit、Gradio、Web或命令行演示界面；",
            "支持选择学生水平、输入问题、查看回答和教材来源；",
            "配置Development和Staging环境；",
            "建立基础CI或自动检查；",
            "编写Smoke Test，检查系统启动、索引加载、Retrieval API、JSON解析和citation ID；",
            "记录常见错误、日志位置和处理方法；",
            "编写README、安装说明和客户演示操作说明；",
            "与Leader共同准备周五Staging Demo。",
        ],
        [
            "环境与依赖配置、.env.example和一键启动脚本；",
            "简单演示界面；",
            "Development和Staging配置；",
            "基础CI或自动检查及Smoke Test；",
            "README、安装说明、操作说明和常见错误处理；",
            "周五可操作的Staging Demo。",
        ],
        [
            "新成员可以按照README启动系统；",
            "客户可以选择水平、输入问题并查看回答与来源；",
            "真实密钥没有进入代码仓库；",
            "Smoke Test只检查可运行性，不输出研究效果结论；",
            "Leader不需要手工拼接模块才能演示。",
        ],
    )

    add_heading(doc, "3. 协作组合与依赖处理", 1)
    add_table(
        doc,
        ["协作组合", "成员", "共同负责"],
        [
            ["教材处理组", "2号 + 4号", "OpenStax解析、Chunking、Metadata和原文追溯"],
            ["索引检索组", "5号 + 6号", "Embedding、FAISS、Top-K Retrieval和Backend API"],
            ["问答生成组", "3号 + 7号", "演示问题、Student Profile、Prompt和LLM"],
            ["集成发布组", "1号 + 8号", "架构、主流水线、运行环境、Staging和客户Demo"],
        ],
        [2100, 1900, 5360],
        font_size=9.4,
    )

    add_heading(doc, "防止人员等待的规则", 2)
    for item in [
        "Leader先冻结接口并提供Fixture数据；",
        "每个模块先交约10%的样例，再扩大到本周目标；",
        "2号先交一个章节，4号即可开始Chunking；",
        "4号先交20–50个chunks，5号即可建立测试索引；",
        "5号先交小型索引，6号即可开发Retriever；",
        "6号未完成时，7号使用固定RetrievalResult开发；",
        "后端未完成时，8号使用Mock数据开发界面；",
        "周五不是第一次集成，主分支在本周内持续保持可运行。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4. 第一周统一验收标准", 1)
    add_table(
        doc,
        ["验收方向", "最低要求"],
        [
            ["OpenStax", "2–3个Physics章节完成可靠解析"],
            ["SciQ", "20–30道演示问题完成标准化"],
            ["Chunking", "500-token chunks可生成并追溯原文"],
            ["索引", "一个embedding和FAISS索引可以重复构建和加载"],
            ["Retrieval", "输入问题后能够返回Top-K证据及来源"],
            ["Profile", "支持Beginner、Intermediate和Advanced三档输入"],
            ["Generation", "至少10–20道问题输出固定JSON"],
            ["Citation", "回答只引用实际检索到的chunk ID"],
            ["Integration", "一个统一入口能够运行完整流程"],
            ["Demo", "客户可以输入问题、选择水平并查看回答和来源"],
            ["Documentation", "架构、接口、环境配置、README和已知问题完整"],
        ],
        [2300, 7060],
        font_size=9.5,
    )

    doc.add_page_break()
    add_heading(doc, "5. 周五客户演示内容", 1)
    demo_items = [
        "当前系统架构和模块状态；",
        "OpenStax原始教材到标准化文本和chunks的过程；",
        "一个问题的Top-K检索结果和教材来源；",
        "Beginner、Intermediate和Advanced三档输入；",
        "LLM固定JSON回答和citation；",
        "简单演示界面；",
        "当前技术限制和已知问题；",
        "第二周开始Evidence Alignment和评估基础设施的计划。",
    ]
    for item in demo_items:
        add_bullet(doc, item)

    add_callout(
        doc,
        "客户沟通边界",
        "第一周交付的是技术链路验证版本，不包含正式评估结果，不能据此判断某个检索方案、embedding模型或LLM效果更好。",
        fill=LIGHT_BLUE,
        accent=NAVY,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
